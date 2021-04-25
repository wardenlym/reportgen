import os
from seleniumwire import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time

output_dir="/output"
if not os.path.isdir(output_dir):
    output_dir = "."
 
GRAFANA_API_KEY="eyJrIjoidGdvYkhyZTkxS1BldkoyVVZKdmQ4UUNiVGhMNlZUbWgiLCJuIjoid2Vla2x5LXJlcG9ydHMiLCJpZCI6MX0="

chrome_options = Options()
chrome_options.add_argument("--headless")
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-dev-shm-usage')
chrome_options.add_argument('Authorization=Bearer eyJrIjoidGdvYkhyZTkxS1BldkoyVVZKdmQ4UUNiVGhMNlZUbWgiLCJuIjoid2Vla2x5LXJlcG9ydHMiLCJpZCI6MX0=')
 
base_url = "http://172.29.40.27:32001/d/sgrwz7XGz/test-weekly-report?orgId=1&refresh=30s"
headers = {
    'Authorization': 'Bearer %s' % GRAFANA_API_KEY
}

#对应的chromedriver的放置目录
driver = webdriver.Chrome(executable_path=("/usr/bin/chromedriver"), options=chrome_options)


def interceptor(request):
    request.headers['Authorization'] = 'Bearer %s' % GRAFANA_API_KEY

driver.request_interceptor = interceptor
driver.get(base_url + "/")
 
start_time=time.time()
print('this is start_time ',start_time)
 
# driver.find_element_by_id("kw").send_keys("admin")
# driver.find_element_by_id("kw").send_keys("admin")
# driver.find_element_by_id("su").click()

time.sleep(5)

print(driver.get_window_size())
driver.set_window_size(1920,1080)
driver.set_window_size(1680,900)
print(driver.get_window_size())

driver.save_screenshot(output_dir+'/screen.png')
 
driver.close()
 
end_time=time.time()
print('this is end_time ',end_time)

driver.quit()


from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Cluster Weekly Report',0)
document.add_heading('集群资源',1)
document.add_heading('本周/上周对比',2)
# document.add_heading(u'三级标题',3)

paragraph = document.add_paragraph('本周集群资源总览')

#增加图片（此处使用相对位置）
document.add_picture(output_dir+'/screen.png',width=Inches(7))

document.save(output_dir+'/weekly-report.docx')
